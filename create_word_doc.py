#!/usr/bin/env python3
"""
创建Word文档的商业计划书
使用python-docx库生成.docx文件
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_business_plan():
    # 创建文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = "钢化膜公司商业计划书（300万预算优化版）"
    doc.core_properties.author = "白月AI助理"
    doc.core_properties.subject = "商业计划书"
    
    # 添加标题
    title = doc.add_heading('钢化膜公司商业计划书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('（300万预算优化版）', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    doc.add_paragraph('编制日期：2026年4月1日')
    doc.add_paragraph('编制人：白月AI助理')
    doc.add_paragraph('')
    
    # 添加目录
    doc.add_heading('目录', 1)
    doc.add_paragraph('一、项目概述')
    doc.add_paragraph('二、市场分析')
    doc.add_paragraph('三、竞争优势')
    doc.add_paragraph('四、团队组建方案')
    doc.add_paragraph('五、300万预算详细分配')
    doc.add_paragraph('六、盈利预测与投资回报')
    doc.add_paragraph('七、实施路线图')
    doc.add_paragraph('八、风险控制')
    doc.add_paragraph('九、分红激励机制')
    doc.add_paragraph('十、总结建议')
    doc.add_paragraph('')
    
    # 一、项目概述
    doc.add_heading('一、项目概述', 1)
    doc.add_heading('1.1 项目背景', 2)
    doc.add_paragraph('基于自有工厂的钢化膜生产优势，计划创立自主品牌，实现自产自销。通过电商+线下双渠道布局，打造中高端钢化膜品牌。')
    
    doc.add_heading('1.2 核心优势', 2)
    doc.add_paragraph('• 自有工厂：成本控制能力强，产品质量可控')
    doc.add_paragraph('• 集团支持：办公场地、供应链资源')
    doc.add_paragraph('• 轻资产运营：云仓物流、聚焦核心平台')
    
    doc.add_heading('1.3 对标企业', 2)
    doc.add_paragraph('• 闪魔（SmartDevil）：成立于2010年，钢化膜线上市场占有率约35%')
    doc.add_paragraph('• 蓝猩、怪兽等知名品牌')
    
    doc.add_heading('1.4 市场定位', 2)
    doc.add_paragraph('中高端钢化膜品牌，主打"工厂直供、品质保证"')
    doc.add_paragraph('')
    
    # 二、市场分析
    doc.add_heading('二、市场分析', 1)
    doc.add_heading('2.1 市场规模', 2)
    doc.add_paragraph('• 中国手机钢化膜市场规模：200-300亿元/年')
    doc.add_paragraph('• 年增长率：15-20%')
    doc.add_paragraph('• 线上渠道占比：约70%')
    doc.add_paragraph('• 线下渠道占比：约30%')
    
    doc.add_heading('2.2 竞争格局', 2)
    doc.add_paragraph('• 头部品牌：闪魔、蓝猩、怪兽、倍思等')
    doc.add_paragraph('• 价格区间：')
    doc.add_paragraph('  - 低端：5-15元/片')
    doc.add_paragraph('  - 中端：15-30元/片')
    doc.add_paragraph('  - 高端：30-100元/片')
    
    doc.add_heading('2.3 渠道分布', 2)
    doc.add_paragraph('• 线上：天猫、京东、抖音、拼多多')
    doc.add_paragraph('• 线下：手机店、3C卖场、数码城')
    doc.add_paragraph('')
    
    # 三、竞争优势
    doc.add_heading('三、竞争优势', 1)
    
    # 添加表格：成本优势对比
    doc.add_heading('3.1 成本优势（相比代工品牌）', 2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '成本项目'
    hdr_cells[1].text = '代工模式'
    hdr_cells[2].text = '自产自销'
    hdr_cells[3].text = '节省金额'
    
    # 数据行
    data = [
        ['生产成本', '4-7元/片', '2.5-5.3元/片', '1.5-1.7元/片'],
        ['代工厂利润', '2-3元/片', '0元/片', '2-3元/片'],
        ['综合成本', '6-10元/片', '2.5-5.3元/片', '3.5-4.7元/片']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('3.2 毛利率对比', 2)
    doc.add_paragraph('• 代工品牌：50-70%')
    doc.add_paragraph('• 自产自销：75-85%')
    
    doc.add_heading('3.3 其他优势', 2)
    doc.add_paragraph('• 质量控制：直接管控生产，次品率降低')
    doc.add_paragraph('• 反应速度：7天内出新款模具')
    doc.add_paragraph('• 定制能力：为客户定制特殊规格')
    doc.add_paragraph('')
    
    # 四、团队组建方案
    doc.add_heading('四、团队组建方案（14人团队）', 1)
    doc.add_heading('4.1 电商团队（8人）', 2)
    doc.add_paragraph('岗位及薪酬结构：')
    doc.add_paragraph('1. 电商运营总监：8,000元底薪 + 5%净利润分红 + 3%股权期权')
    doc.add_paragraph('2. 平台店长（天猫/抖音）：6,000元底薪 + 销售额1%提成 × 2人')
    doc.add_paragraph('3. 抖音直播运营：6,000元底薪 + 直播销售额3%提成')
    doc.add_paragraph('4. 美工设计师：5,000元底薪 + 设计项目奖金')
    doc.add_paragraph('5. 客服专员：4,000元底薪 + 销售额0.5%提成 × 2人')
    doc.add_paragraph('6. 仓储协调员：3,500元固定薪资 × 2人')
    
    doc.add_heading('4.2 线下销售团队（6人）', 2)
    doc.add_paragraph('1. 销售总监：8,000元底薪 + 团队销售额3%提成 + 2%股权期权')
    doc.add_paragraph('2. 区域销售经理：5,000元底薪 + 个人销售额5%提成 × 2人')
    doc.add_paragraph('3. 渠道专员：4,000元底薪 + 开发客户销售额2%提成 × 2人')
    doc.add_paragraph('4. 市场专员：5,000元底薪 + 活动效果奖金')
    
    doc.add_heading('4.3 年度人力成本', 2)
    doc.add_paragraph('• 电商团队：54-60万元')
    doc.add_paragraph('• 销售团队：42-48万元')
    doc.add_paragraph('• 总计：96-108万元')
    doc.add_paragraph('')
    
    # 五、300万预算详细分配
    doc.add_heading('五、300万预算详细分配', 1)
    
    # 预算表格
    doc.add_heading('5.4 预算汇总表', 2)
    budget_table = doc.add_table(rows=9, cols=4)
    budget_table.style = 'Light Grid'
    
    budget_header = budget_table.rows[0].cells
    budget_header[0].text = '项目'
    budget_header[1].text = '预算（万元）'
    budget_header[2].text = '占比'
    budget_header[3].text = '说明'
    
    budget_data = [
        ['团队成本', '96-108', '32-36%', '14人团队'],
        ['办公分摊', '12', '4%', '集团支持'],
        ['产品开发', '75', '25%', '核心投入'],
        ['平台费用', '20', '6.7%', '天猫+抖音'],
        ['营销推广', '48', '16%', '集中投放'],
        ['云仓物流', '1.2', '0.4%', '2元/平方/天'],
        ['其他费用', '39.8', '13.3%', '杂项及备用'],
        ['总计', '300', '100%', '']
    ]
    
    for i, row_data in enumerate(budget_data, 1):
        row_cells = budget_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph('')
    
    # 六、盈利预测
    doc.add_heading('六、盈利预测与投资回报', 1)
    doc.add_heading('6.1 保守预测', 2)
    doc.add_paragraph('• 年销售额：804万元')
    doc.add_paragraph('• 毛利率：75-80%')
    doc.add_paragraph('• 毛利润：603-643万元')
    doc.add_paragraph('• 运营费用：240万元')
    doc.add_paragraph('• 净利润：363-403万元')
    doc.add_paragraph('• 净利率：45-50%')
    doc.add_paragraph('• 投资回收期：0.7-0.8年')
    
    doc.add_heading('6.2 乐观预测', 2)
    doc.add_paragraph('• 年销售额：1,275万元')
    doc.add_paragraph('• 毛利率：78-83%')
    doc.add_paragraph('• 毛利润：994-1,058万元')
    doc.add_paragraph('• 运营费用：280万元')
    doc.add_paragraph('• 净利润：714-778万元')
    doc.add_paragraph('• 净利率：56-61%')
    doc.add_paragraph('• 投资回收期：0.4-0.5年')
    doc.add_paragraph('')
    
    # 七、实施路线图
    doc.add_heading('七、实施路线图', 1)
    doc.add_heading('7.1 第1个月', 2)
    doc.add_paragraph('• 完成公司注册')
    doc.add_paragraph('• 核心团队到位（5人）')
    doc.add_paragraph('• 开发首批10款模具')
    doc.add_paragraph('• 申请天猫店铺')
    
    doc.add_heading('7.2 第2个月', 2)
    doc.add_paragraph('• 天猫店铺上线')
    doc.add_paragraph('• 首批产品生产完成')
    doc.add_paragraph('• 申请抖音店铺')
    doc.add_paragraph('• 启动小规模测试')
    
    doc.add_heading('7.3 第3个月', 2)
    doc.add_paragraph('• 抖音店铺上线')
    doc.add_paragraph('• 产品线扩大到20款')
    doc.add_paragraph('• 启动正式营销')
    doc.add_paragraph('• 建立基础销量')
    doc.add_paragraph('')
    
    # 八、风险控制
    doc.add_heading('八、风险控制', 1)
    doc.add_heading('8.1 资金风险控制', 2)
    doc.add_paragraph('• 每月现金流分析')
    doc.add_paragraph('• 安全线：账户余额不低于50万')
    doc.add_paragraph('• 预备融资渠道')
    
    doc.add_heading('8.2 团队风险控制', 2)
    doc.add_paragraph('• 核心团队签订3年服务协议')
    doc.add_paragraph('• 股权分期兑现（4年）')
    doc.add_paragraph('• 设置竞业禁止条款')
    doc.add_paragraph('')
    
    # 九、分红机制
    doc.add_heading('九、分红激励机制', 1)
    doc.add_heading('9.1 利润分配比例', 2)
    doc.add_paragraph('年度净利润分配：')
    doc.add_paragraph('1. 创始人：60%')
    doc.add_paragraph('2. 核心团队分红池：25%')
    doc.add_paragraph('   • 电商总监：10%')
    doc.add_paragraph('   • 销售总监：7.5%')
    doc.add_paragraph('   • 其他核心：7.5%')
    doc.add_paragraph('3. 员工奖金池：10%')
    doc.add_paragraph('4. 发展基金：5%')
    doc.add_paragraph('')
    
    # 十、总结
    doc.add_heading('十、总结建议', 1)
    doc.add_heading('10.1 项目亮点', 2)
    doc.add_paragraph('1. 自有工厂成本优势明显')
    doc.add_paragraph('2. 300万预算经过多轮优化')
    doc.add_paragraph('3. 轻资产运营模式')
    doc.add_paragraph('4. 激励性薪酬结构')
    doc.add_paragraph('5. 快速盈利预期')
    
    doc.add_heading('10.2 成功关键', 2)
    doc.add_paragraph('1. 利用集团资源降低固定成本')
    doc.add_paragraph('2. 聚焦天猫+抖音核心平台')
    doc.add_paragraph('3. 快速验证、小步快跑')
    doc.add_paragraph('4. 团队利益与公司深度绑定')
    
    doc.add_heading('10.3 下一步建议', 2)
    doc.add_paragraph('1. 立即启动公司注册')
    doc.add_paragraph('2. 招募核心团队成员')
    doc.add_paragraph('3. 同步开发首批产品')
    doc.add_paragraph('4. 准备平台入驻材料')
    
    # 保存文档
    output_path = "/root/.openclaw/workspace/钢化膜公司商业计划书_正式版.docx"
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        # 检查是否安装了python-docx
        import docx
        file_path = create_business_plan()
        print(f"✅ Word文档创建成功: {file_path}")
    except ImportError:
        print("❌ 需要安装python-docx库")
        print("请运行: pip install python-docx")
        exit(1)